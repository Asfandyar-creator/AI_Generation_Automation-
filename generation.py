import os
import pandas as pd
import openai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import threading
from dotenv import load_dotenv
import logging
import time
import re
import json
import sys

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    stream=open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1),
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Load environment variables
load_dotenv()

# Configure OpenAI API - Updated for v1.x
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    logging.error("OPENAI_API_KEY not found in environment variables or .env file")
else:
    logging.info("API key loaded successfully")
    client = openai.OpenAI(api_key=api_key)

# Global processing status
processing_status = {"running": False, "current_file": "", "completed": False}


def load_processed_clusters(checkpoint_file="processed_clusters.json"):
    """
    Load the list of already processed clusters from a JSON file
    """
    if os.path.exists(checkpoint_file):
        try:
            with open(checkpoint_file, "r") as f:
                return json.load(f)
        except Exception as e:
            logging.error(f"Error loading processed clusters file: {str(e)}")
            return []
    return []

def save_processed_cluster(cluster, checkpoint_file="processed_clusters.json"):
    """
    Save a cluster to the processed clusters JSON file
    """
    processed_clusters = load_processed_clusters(checkpoint_file)
    if cluster not in processed_clusters:
        processed_clusters.append(cluster)
        try:
            with open(checkpoint_file, "w") as f:
                json.dump(processed_clusters, f)
            logging.info(f"Saved cluster '{cluster}' to processed clusters file")
        except Exception as e:
            logging.error(f"Error saving to processed clusters file: {str(e)}")

def clean_text(text):
    """
    Clean text by removing unnecessary quotes and standardizing formatting
    """
    # Remove various types of quotes that might cause issues
    text = text.replace('"', '').replace('"', '').replace('"', '')
    # Replace fancy apostrophes with standard ones
    text = text.replace("'", "'")
    return text

def process_urls_and_organizations(urls_text):
    """
    Process multiple URLs from a text field and extract organization names
    """
    if not urls_text or not isinstance(urls_text, str):
        return []
    
    results = []
    
    # Split by common separators (newline, comma, semicolon)
    urls = re.split(r'[\n,;]+', urls_text)
    
    for url in urls:
        url = url.strip()
        if not url:
            continue
            
        # Extract organization name
        org_name = extract_organization_from_url(url)
        results.append({"url": url, "org_name": org_name})
    
    return results

def extract_organization_from_url(url):
    """
    Extract organization name from URL
    """
    if not url or not isinstance(url, str):
        return ""
    
    try:
        # Remove http:// or https:// and www.
        clean_url = url.lower().replace("https://", "").replace("http://", "").replace("www.", "")
        
        # Get domain parts
        domain_parts = clean_url.split("/")[0].split(".")
        
        # Most likely the organization name is the first part of the domain
        org_name = domain_parts[0]
        
        # Convert to title case and replace hyphens with spaces
        org_name = org_name.replace("-", " ").replace("_", " ").title()
        
        return org_name
    except Exception as e:
        logging.error(f"Error extracting organization name from URL: {str(e)}")
        return ""   

def test_api_connection():
    """Test connection to OpenAI API"""
    try:
        logging.info("Testing OpenAI API connection...")
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": "Hello, this is a test."}],
            max_tokens=10
        )
        logging.info("API connection successful")
        return True
    except Exception as e:
        logging.error(f"API connection failed: {str(e)}")
        return False, str(e)

def generate_seo_title(original_title):
    """
    Generates an SEO optimized title based on the original title
    """
    logging.info(f"Generating SEO title for: {original_title}")
    try:
        prompt = f"""
            
            """


        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=100
        )

        seo_title = response.choices[0].message.content.strip()
        seo_title = clean_text(seo_title)
        logging.info(f"SEO title generated: {seo_title}")
        return seo_title
    except Exception as e:
        logging.error(f"Error generating SEO title: {str(e)}")
        # Fallback to original title in case of error
        return original_title

def generate_duration_and_level(title, context, programme):
    """
    Generates training duration and level based on the title, context and programme
    """
    logging.info(f"Generating duration and level for: {title}")
    try:
        prompt = f"""
           
            """

        # Removed the response_format parameter that was causing the error
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=200
        )

        content = response.choices[0].message.content.strip()
        
        # More robust JSON parsing approach
        import json
        import re
        
        # Try to extract just the JSON part if there's any extra text
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            json_str = json_match.group(0)
        else:
            json_str = content
            
        try:
            # Try standard JSON parsing
            data = json.loads(json_str)
        except json.JSONDecodeError:
            # If that fails, try the eval approach with proper safeguards
            sanitized = json_str.replace('null', 'None').replace('true', 'True').replace('false', 'False')
            data = eval(sanitized)
        
        # Validate and sanitize the data
        if not isinstance(data, dict):
            raise ValueError("Response is not a dictionary")
        
        # Check required keys
        required_keys = ["duree_heures", "duree_jours", "niveau"]
        for key in required_keys:
            if key not in data:
                raise ValueError(f"Missing required key: {key}")
        
        # Ensure values are of correct type
        try:
            data["duree_heures"] = int(float(data["duree_heures"]))
            data["duree_jours"] = float(data["duree_jours"])
        except (ValueError, TypeError):
            raise ValueError("Invalid duration values")
        
        # Ensure niveau is one of the allowed values
        allowed_levels = ["Initiation", "Avancé", "Expert"]
        if data["niveau"] not in allowed_levels:
            # Try to find the closest match
            for level in allowed_levels:
                if level.lower() in data["niveau"].lower():
                    data["niveau"] = level
                    break
            else:
                data["niveau"] = "Initiation"  # Default if no match
        
        logging.info(f"Duration and level generated: {data}")
        return data
    except Exception as e:
        logging.error(f"Error generating duration and level: {str(e)}")
        
        # Smart fallback logic based on content analysis
        try:
            # Estimate based on programme complexity
            programme_length = len(programme.strip()) if programme else 0
            
            # Analyze the title and content for complexity indicators
            combined_text = (title + " " + (context or "") + " " + (programme or "")).lower()
            
            # Default values
            hours = 14
            days = 2
            level = "Initiation"
            
            # Adjust based on text length
            if programme_length > 3000:
                hours = 28
                days = 4
                level = "Avancé"
            elif programme_length > 1500:
                hours = 21
                days = 3
                level = "Avancé"
            elif programme_length < 500:
                hours = 7
                days = 1
            
            # Adjust based on keywords
            advanced_keywords = ["avancé", "perfectionnement", "approfondi", "spécialisé", "intermédiaire"]
            expert_keywords = ["expert", "mastérisation", "spécialisation", "maîtrise", "professionnel"]
            
            if any(word in combined_text for word in expert_keywords):
                level = "Expert"
                if hours < 21:  # Ensure Expert courses are at least 3 days
                    hours = 21
                    days = 3
            elif any(word in combined_text for word in advanced_keywords):
                level = "Avancé"
                if hours < 14:  # Ensure Advanced courses are at least 2 days
                    hours = 14
                    days = 2
            
            # Check for specific training types that typically have standard durations
            if "certifi" in combined_text or "diplôme" in combined_text:
                hours = 35
                days = 5
                level = "Expert"
            
            fallback_data = {
                "duree_heures": hours,
                "duree_jours": days,
                "niveau": level
            }
            
            logging.info(f"Using smart fallback values based on content analysis: {fallback_data}")
            return fallback_data
            
        except Exception as inner_e:
            logging.error(f"Error in fallback logic: {str(inner_e)}")
            
            # Ultimate fallback with variety instead of fixed values
            import random
            hours_options = [7, 14, 21, 28, 35]
            days_options = [1, 2, 3, 4, 5]  # Corresponding to hours
            level_options = ["Initiation", "Avancé", "Expert"] 
            
            # Choose random but ensure they match (hours = days * 7)
            index = random.randint(0, 4)
            hours = hours_options[index]
            days = days_options[index]
            level = random.choice(level_options)
            
            logging.warning(f"Using random fallback values: {hours}h, {days}d, {level}")
            return {"duree_heures": hours, "duree_jours": days, "niveau": level}

def generate_context(title, context, target_audience, prerequisites, objectives, programme):
    """
    Generates SEO optimized context for the training
    """
    logging.info(f"Generating context for: {title}")
    try:
        prompt = f"""
            
            """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=4000
        )

        content = response.choices[0].message.content.strip()
        content = clean_text(content)
        logging.info(f"Context generated successfully for {title} ({len(content)} chars)")
        return content
    except Exception as e:
        logging.error(f"Error generating context: {str(e)}")
        raise

def generate_target_audience(title, target_audience):
    """
    Generates the "À qui s'adresse cette formation ?" section
    """
    logging.info(f"Generating target audience for: {title}")
    try:
        prompt = f"""
            
            """
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=500
        )

        content = response.choices[0].message.content.strip()
        content = clean_text(content)
        logging.info(f"Target audience generated successfully for {title}")
        return content
    except Exception as e:
        logging.error(f"Error generating target audience: {str(e)}")
        raise

def generate_prerequisites(title, prerequisites):
    """
    Generates the "Prérequis" section
    """
    logging.info(f"Generating prerequisites for: {title}")
    try:
        prompt = f"""
            
            """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=500
        )

        content = response.choices[0].message.content.strip()
        content = clean_text(content)
        logging.info(f"Prerequisites generated successfully for {title}")
        return content
    except Exception as e:
        logging.error(f"Error generating prerequisites: {str(e)}")
        raise

def generate_objectives(title, objectives):
    """
    Generates the "Objectifs" section
    """
    logging.info(f"Generating objectives for: {title}")
    try:
        prompt = f"""
            
    """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=500
        )

        content = response.choices[0].message.content.strip()
        content = clean_text(content)
        logging.info(f"Objectives generated successfully for {title}")
        return content
    except Exception as e:
        logging.error(f"Error generating objectives: {str(e)}")
        raise

def generate_program(title, programme):
    """
    Generates the "Programme" section
    """
    logging.info(f"Generating program for: {title}")
    try:
        prompt = f"""
            
            """

        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )

        content = response.choices[0].message.content.strip()
        content = clean_text(content)
        logging.info(f"Program generated successfully for {title}")
        return content
    except Exception as e:
        logging.error(f"Error generating program: {str(e)}")
        raise

import re
import logging

def parse_competitors_prices(competitors_prices_text):
    """
    Parse competitors prices data from the specified format with improved robustness.
    Handles varied formats and special cases better.
    """
    if not competitors_prices_text or not isinstance(competitors_prices_text, str):
        return []
        
    result = []
    competitors = competitors_prices_text.split("|")
    
    for idx, competitor in enumerate(competitors, 1):
        competitor = competitor.strip()
        if not competitor:
            continue
            
        try:
            # Look for the organization-->title pattern
            if "--->" not in competitor:
                logging.warning(f"Separator '-->' not found in competitor: {competitor}")
                continue
                
            # Split at the first occurrence of --->
            parts = competitor.split("--->", 1)
            if len(parts) != 2:
                logging.warning(f"Invalid competitor format: {competitor}")
                continue
                
            org_name = parts[0].strip()
            remaining_text = parts[1].strip()
            
            # Find the last -- which should separate title from prices
            # This handles cases where -- appears in the title itself
            if "--(" not in remaining_text:
                # Try a more forgiving approach - look for price pattern
                match = re.search(r'(.*?)(\(\S.*?\))', remaining_text)
                if match:
                    training_title = match.group(1).strip()
                    prices_part = remaining_text[match.start(2):]
                else:
                    logging.warning(f"Cannot identify title and prices separation in: {remaining_text}")
                    continue
            else:
                title_prices_split = remaining_text.split("--(", 1)
                training_title = title_prices_split[0].strip()
                prices_part = "(" + title_prices_split[1]
            
            # Extract all groups within parentheses
            prices_matches = re.findall(r"\((.*?)\)", prices_part)

            # Ensure we handle missing prices and additional data
            individual_price = prices_matches[0].strip() if len(prices_matches) > 0 else ""
            inter_company_price = prices_matches[1].strip() if len(prices_matches) > 1 else ""
            intra_company_price = prices_matches[2].strip() if len(prices_matches) > 2 else ""
            url = prices_matches[3].strip() if len(prices_matches) > 3 else ""
            duration = prices_matches[4].strip() if len(prices_matches) > 4 else ""
            level = prices_matches[5].strip() if len(prices_matches) > 5 else ""

            result.append({
                "sr_no": idx,
                "organization": org_name,
                "training_title": training_title,
                "individual_price": individual_price,
                "inter_company_price": inter_company_price,
                "intra_company_price": intra_company_price,
                "url": url,
                "duration": duration,
                "level": level
            })
        except Exception as e:
            logging.error(f"Error parsing competitor data: {str(e)} - {competitor}")
    
    return result


def create_docx(output_path, title, pricing_data, context_content, target_audience_content, 
               prerequisites_content, objectives_content, programme_content, url_org_data=None,
               competitors_data=None):
    """
    Creates a Word document with the generated content
    """
    logging.info(f"Creating Word document at: {output_path}")
    try:
        doc = Document()
        
        # Clean all text inputs
        title = clean_text(title)
        context_content = clean_text(context_content)
        target_audience_content = clean_text(target_audience_content)
        prerequisites_content = clean_text(prerequisites_content)
        objectives_content = clean_text(objectives_content)
        programme_content = clean_text(programme_content)
        
        # Add the title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add pricing info section
        pricing_para = doc.add_paragraph()
        pricing_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Format duration
        duree_text = f"Durée : {pricing_data['duree_heures']} heures soit {pricing_data['duree_jours_display']}"  
        duration_run = pricing_para.add_run(duree_text)
        duration_run.bold = True
        duration_run.font.size = Pt(12)
        
        # Add a line break
        pricing_para.add_run("\n")
        
        # Format level
        level_text = f"Niveau de formation : {pricing_data['niveau']}"
        level_run = pricing_para.add_run(level_text)
        level_run.bold = True
        level_run.font.size = Pt(12)

        # Add cluster if available
        if 'cluster' in pricing_data and pricing_data['cluster']:
            pricing_para.add_run("\n")
            cluster_text = f"Cluster : {pricing_data['cluster']}"
            cluster_run = pricing_para.add_run(cluster_text)
            cluster_run.bold = True
            cluster_run.font.size = Pt(12)
        
        # Add pricing details
        pricing_para.add_run("\n\n")
        pricing_info = pricing_para.add_run("Tarifs :")
        pricing_info.bold = True
        pricing_info.font.size = Pt(12)
        
        # pricing_para.add_run("\n")
        # prix_inter_text = f"Prix inter-entreprise HT : {pricing_data['prix_inter']}"
        # pricing_para.add_run(prix_inter_text).font.size = Pt(11)
        
        pricing_para.add_run("\n")
        prix_individuel_text = f"Prix individuel HT : {pricing_data['prix_individuel']}"
        pricing_para.add_run(prix_individuel_text).font.size = Pt(11)
        
        # pricing_para.add_run("\n")
        # prix_intra_text = f"Prix intra-entreprise HT : {pricing_data['prix_intra']}"
        # pricing_para.add_run(prix_intra_text).font.size = Pt(11)
        
        pricing_para.add_run("\n")
        prix_achat_text = f"Prix d'achat HT : {pricing_data['prix_achat']}"
        pricing_para.add_run(prix_achat_text).font.size = Pt(11)
        
        pricing_para.add_run("\n")
        prix_vente_text = f"Prix de vente interne HT : {pricing_data['prix_vente']}"
        pricing_para.add_run(prix_vente_text).font.size = Pt(11)
        
        # Context section
        doc.add_heading("Contexte", level=1)
        doc.add_paragraph(context_content)
        
        # Target audience section
        doc.add_heading("À qui s'adresse cette formation ?", level=1)
        doc.add_paragraph(target_audience_content)
        
        # Prerequisites section
        doc.add_heading("Prérequis", level=1)
        # Process bullet points
        for line in prerequisites_content.split("\n"):
            if line.strip().startswith("•") or line.strip().startswith("-"):
                doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
            elif line.strip():
                doc.add_paragraph(line.strip())
        
        # Objectives section
        doc.add_heading("Objectifs", level=1)
        # Process bullet points
        for line in objectives_content.split("\n"):
            if line.strip().startswith("•") or line.strip().startswith("-"):
                doc.add_paragraph(line.strip()[1:].strip(), style='List Bullet')
            elif line.strip():
                doc.add_paragraph(line.strip())
        
        # Program section
        doc.add_heading("Programme", level=1)
        current_chapter = None
        
        for line in programme_content.split("\n"):
            line = line.strip()
            if not line:
                continue
                
            # Detect chapter titles (uppercase or starting with #)
            if line.isupper() or (line.startswith("CHAPITRE") or line.startswith("#")):
                current_chapter = doc.add_paragraph()
                chapter_run = current_chapter.add_run(line)
                chapter_run.bold = True
            # Detect bullet points
            elif line.startswith("•") or line.strip().startswith("-"):
                doc.add_paragraph(line[1:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        # Add competitors pricing table if data exists
        # Add competitors pricing table if data exists
        if competitors_data and len(competitors_data) > 0:
            doc.add_heading("Analyse comparative des prix des concurrents", level=1)
            
            table_description = doc.add_paragraph("Tableau comparatif des prix pratiqués par les concurrents pour des formations similaires.")
            
            # Create the table with headers
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            headers = ["Sr. No.", "Organisation", "Titre de formation", "Prix individuel", "Prix inter-entreprise", "Prix intra-entreprise", "URL", "Durée", "Niveau"]
            
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                # Bold formatting for header row
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add competitor data rows
            for competitor in competitors_data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(competitor["sr_no"])
                row_cells[1].text = competitor["organization"]
                row_cells[2].text = competitor["training_title"]
                row_cells[3].text = competitor["individual_price"]
                row_cells[4].text = competitor["inter_company_price"]
                row_cells[5].text = competitor["intra_company_price"]
                row_cells[6].text = competitor.get("url", "")
                row_cells[7].text = competitor.get("duration", "")
                row_cells[8].text = competitor.get("level", "")


        # Save the document
        doc.save(output_path)
        logging.info(f"Document saved successfully at: {output_path}")
        return True
    except Exception as e:
        logging.error(f"Error creating Word document: {str(e)}")
        raise

def sanitize_filename(filename):
    """
    Sanitize a string to be used as a filename
    """
    # Replace invalid characters with underscores
    sanitized = re.sub(r'[\\/*?:"<>|]', "_", filename)
    # Remove any leading/trailing whitespace
    sanitized = sanitized.strip()
    # Limit length to avoid very long filenames
    if len(sanitized) > 150:
        sanitized = sanitized[:150]
    return sanitized

def process_excel_file(file_path, output_folder, skip_processed_clusters=True):
    """
    Process an Excel file and generate a Word document
    """
    logging.info(f"Processing file: {file_path}")
    try:
        # Read the Excel file
        df = pd.read_excel(file_path)
        
        # NEW CODE: Check if this file has a cluster that's already been processed
        cluster = df["Cluster"].iloc[0] if "Cluster" in df.columns and not pd.isna(df["Cluster"].iloc[0]) else ""
        
        # NEW CODE: Skip processing if cluster has already been processed
        if skip_processed_clusters and cluster:
            processed_clusters = load_processed_clusters()
            if cluster in processed_clusters:
                skip_message = f"Skipping {file_path} - Cluster '{cluster}' already processed"
                logging.info(skip_message)
                return skip_message
        
        # Check required columns
        required_columns = ["title", "context", "target_audience", "prerequisites", "objectives", "programme"]
        for col in required_columns:
            if col not in df.columns:
                error_message = f"Column '{col}' missing in {file_path}"
                logging.error(error_message)
                return error_message
        
        # Get the data
        title = df["title"].iloc[0]
        context = df["context"].iloc[0] if not pd.isna(df["context"].iloc[0]) else ""
        target_audience = df["target_audience"].iloc[0] if not pd.isna(df["target_audience"].iloc[0]) else ""
        prerequisites = df["prerequisites"].iloc[0] if not pd.isna(df["prerequisites"].iloc[0]) else ""
        objectives = df["objectives"].iloc[0] if not pd.isna(df["objectives"].iloc[0]) else ""
        programme = df["programme"].iloc[0] if not pd.isna(df["programme"].iloc[0]) else ""
        urls_text = df["url"].iloc[0] if "url" in df.columns and not pd.isna(df["url"].iloc[0]) else ""
        cluster = df["Cluster"].iloc[0] if "Cluster" in df.columns and not pd.isna(df["Cluster"].iloc[0]) else ""
        
        # Get competitors prices data
        competitors_prices_text = df["competitors_prices"].iloc[0] if "competitors_prices" in df.columns and not pd.isna(df["competitors_prices"].iloc[0]) else ""
        competitors_data = parse_competitors_prices(competitors_prices_text)
        
        # Get pricing data - extract numeric parts from text fields when necessary
        def extract_numeric(value):
            if pd.isna(value):
                return ""
            # Extract only numeric characters from the string
            import re
            numbers = re.findall(r'\d+', str(value))
            if numbers:
                return int(numbers[0])
            return ""
        
        duree_jours_raw = df["Durée (jours)"].iloc[0] if "Durée (jours)" in df.columns and not pd.isna(df["Durée (jours)"].iloc[0]) else ""
        duree_jours = extract_numeric(duree_jours_raw)
        duree_jours_display = str(duree_jours_raw)  # Keep original text for display
        
        prix_inter = df["Prix inter-entreprise HT"].iloc[0] if "Prix inter-entreprise HT" in df.columns and not pd.isna(df["Prix inter-entreprise HT"].iloc[0]) else ""
        prix_individuel = df["Prix individuel HT"].iloc[0] if "Prix individuel HT" in df.columns and not pd.isna(df["Prix individuel HT"].iloc[0]) else ""
        prix_intra = df["Prix intra-entreprise HT"].iloc[0] if "Prix intra-entreprise HT" in df.columns and not pd.isna(df["Prix intra-entreprise HT"].iloc[0]) else ""
        prix_achat = df["Prix d'achat HT"].iloc[0] if "Prix d'achat HT" in df.columns and not pd.isna(df["Prix d'achat HT"].iloc[0]) else ""
        prix_vente = df["Prix de vente interne HT"].iloc[0] if "Prix de vente interne HT" in df.columns and not pd.isna(df["Prix de vente interne HT"].iloc[0]) else ""
        
        # Process URLs and extract organization names
        url_org_data = process_urls_and_organizations(urls_text)
        
        logging.info(f"Starting content generation for: {title}")
        
        # Generate SEO title
        seo_title = generate_seo_title(title)
        
        # Only generate level, use provided duration
        level = generate_duration_and_level(title, context, programme)["niveau"]
        
        # Create a pricing data dictionary
        pricing_data = {
            "duree_jours": duree_jours,
            "duree_jours_display": duree_jours_display,
            "duree_heures": int(duree_jours) * 7 if duree_jours else 0,  # Calculate hours based on days
            "niveau": level,
            "cluster": cluster,
            "prix_inter": prix_inter,
            "prix_individuel": prix_individuel,
            "prix_intra": prix_intra,
            "prix_achat": prix_achat,
            "prix_vente": prix_vente
        }
        
        # Generate content
        context_content = generate_context(title, context, target_audience, prerequisites, objectives, programme)
        target_audience_content = generate_target_audience(title, target_audience)
        prerequisites_content = generate_prerequisites(title, prerequisites)
        objectives_content = generate_objectives(title, objectives)
        programme_content = generate_program(title, programme)
        
        # Create the docx file with SEO title as filename
        sanitized_seo_title = sanitize_filename(seo_title)
        output_path = os.path.join(output_folder, f"{sanitized_seo_title}.docx")
        
        create_docx(output_path, seo_title, pricing_data, context_content, target_audience_content, 
           prerequisites_content, objectives_content, programme_content, url_org_data, competitors_data)
        
        # NEW CODE: After successful processing, save this cluster as processed
        if cluster:
            save_processed_cluster(cluster)
        
        success_message = f"Processing successful for {file_path} - Saved as: {sanitized_seo_title}.docx"
        logging.info(success_message)
        return success_message
        
    except openai.OpenAIError as e:
        error_message = f"OpenAI API Error: {str(e)}"
        logging.error(error_message)
        return error_message
    except Exception as e:
        error_message = f"Error processing {file_path}: {str(e)}"
        logging.error(error_message)
        return error_message
        
        
def process_folder(input_folder, output_folder, skip_processed=True):
    """
    Process all Excel files in a folder
    """
    global processing_status
    processing_status["running"] = True
    processing_status["completed"] = False
    
    # Create output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        logging.info(f"Created output folder: {output_folder}")
    
    # Test API connection before processing
    api_test = test_api_connection()
    if isinstance(api_test, tuple) and not api_test[0]:
        error_message = f"API connection test failed: {api_test[1]}"
        logging.error(error_message)
        messagebox.showerror("API Error", f"Connection to OpenAI API failed: {api_test[1]}\n\nPlease check your API key and try again.")
        processing_status["running"] = False
        return
    
    # Process each Excel file
    excel_files = [f for f in os.listdir(input_folder) if f.endswith(".xlsx")]
    
    if not excel_files:
        logging.warning(f"No Excel files found in {input_folder}")
        messagebox.showinfo("No Files", f"No Excel files found in the selected folder.")
        processing_status["running"] = False
        return
    
    for file in excel_files:
        file_path = os.path.join(input_folder, file)
        processing_status["current_file"] = file
        result = process_excel_file(file_path, output_folder, skip_processed)
        log_widget.insert(tk.END, f"{result}\n")
        log_widget.see(tk.END)
    
    processing_status["running"] = False
    processing_status["completed"] = True
    messagebox.showinfo("Processing Complete", "All files have been processed!")

def select_input_folder():
    """
    Select input folder
    """
    folder_selected = filedialog.askdirectory()
    if folder_selected:
        input_folder_var.set(folder_selected)
        logging.info(f"Input folder selected: {folder_selected}")

def select_output_folder():
    """
    Select output folder
    """
    folder_selected = filedialog.askdirectory()
    if folder_selected:
        output_folder_var.set(folder_selected)
        logging.info(f"Output folder selected: {folder_selected}")

def run_processing():
    """
    Start processing files
    """
    input_folder = input_folder_var.get()
    output_folder = output_folder_var.get()
    skip_processed = skip_processed_var.get()
    
    if not input_folder:
        messagebox.showerror("Error", "Please select an input folder.")
        return
        
    if not output_folder:
        messagebox.showerror("Error", "Please select an output folder.")
        return
    
    # Clear log widget
    log_widget.delete(1.0, tk.END)
    
    # Check for API key before starting
    if not api_key:
        messagebox.showerror("API Key Missing", "OpenAI API key not found. Please add it to your .env file.")
        logging.error("API key missing or invalid")
        return
    
    threading.Thread(target=process_folder, args=(input_folder, output_folder, skip_processed)).start()
    status_label.config(text="Processing in progress...")
    root.after(1000, check_status)

def check_status():
    """
    Check processing status
    """
    if processing_status["running"]:
        status_label.config(text=f"Processing: {processing_status['current_file']}")
        root.after(1000, check_status)
    elif processing_status["completed"]:
        status_label.config(text="Processing Complete!")

# Tkinter Interface
root = tk.Tk()
root.title("Excel to Word Converter")
root.geometry("650x500")

# Folder path variables
input_folder_var = tk.StringVar()
output_folder_var = tk.StringVar()

# Main frame
main_frame = tk.Frame(root, padx=20, pady=20)
main_frame.pack(fill=tk.BOTH, expand=True)

# Input folder selection
input_frame = tk.Frame(main_frame)
input_frame.pack(fill=tk.X, pady=5)
tk.Label(input_frame, text="Excel Files Folder:").pack(side=tk.LEFT)
tk.Entry(input_frame, textvariable=input_folder_var, width=50).pack(side=tk.LEFT, padx=5)
tk.Button(input_frame, text="Browse", command=select_input_folder).pack(side=tk.LEFT)

# Output folder selection
output_frame = tk.Frame(main_frame)
output_frame.pack(fill=tk.X, pady=5)
tk.Label(output_frame, text="Word Output Folder:").pack(side=tk.LEFT)
tk.Entry(output_frame, textvariable=output_folder_var, width=50).pack(side=tk.LEFT, padx=5)
tk.Button(output_frame, text="Browse", command=select_output_folder).pack(side=tk.LEFT)


skip_processed_var = tk.BooleanVar(value=True)
skip_frame = tk.Frame(main_frame)
skip_frame.pack(fill=tk.X, pady=5)
tk.Checkbutton(skip_frame, text="Skip already processed clusters", variable=skip_processed_var).pack(side=tk.LEFT)

# Start button
tk.Button(main_frame, text="Convert Files", command=run_processing, 
          bg="#4CAF50", fg="white", font=("Arial", 12), padx=10, pady=5).pack(pady=15)

# Status label
status_label = tk.Label(main_frame, text="", font=("Arial", 10))
status_label.pack(pady=5)

# Log widget
log_frame = tk.Frame(main_frame)
log_frame.pack(fill=tk.BOTH, expand=True, pady=10)
tk.Label(log_frame, text="Processing Log:").pack(anchor=tk.W)
log_widget = scrolledtext.ScrolledText(log_frame, height=10)
log_widget.pack(fill=tk.BOTH, expand=True)

# Info label
info_label = tk.Label(main_frame, text="This program converts Excel files to SEO-optimized Word documents.", 
                     font=("Arial", 8), fg="gray")
info_label.pack(side=tk.BOTTOM, pady=10)

# Check API key on startup
if not api_key:
    log_widget.insert(tk.END, "WARNING: OpenAI API key not found. Please add it to your .env file.\n")
    log_widget.see(tk.END)

root.mainloop()